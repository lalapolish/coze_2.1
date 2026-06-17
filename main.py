import re
import io
import os
import uuid
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from fastapi import FastAPI, Request
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

app = FastAPI(openapi_version="3.0.0")

if not os.path.exists("static"):
    os.makedirs("static")
app.mount("/static", StaticFiles(directory="static"), name="static")

class ReportInput(BaseModel):
    ch1_text: Optional[str] = ""  
    ch2_text: Optional[str] = ""
    ch3_text: Optional[str] = ""
    ch4_text: Optional[str] = ""
    ch5_text: Optional[str] = ""
    ch6_text: Optional[str] = ""
    ch7_text: Optional[str] = ""
    ch8_text: Optional[str] = "" 
    appendix: Optional[str] = "" # 新增附录变量

# ================= 解决图片方框：强制加载本地 SimHei.ttf =================
font_path = os.path.join(os.getcwd(), 'SimHei.ttf')
if os.path.exists(font_path):
    fm.fontManager.addfont(font_path)
    plt.rcParams['font.sans-serif'] = ['SimHei']
else:
    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

def set_font(run, size=12, bold=False, color=None):
    """设置字体、大小、加粗及颜色"""
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)

def set_cell_shading(cell, color):
    """设置单元格背景颜色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_table_border(table):
    """为表格设置蓝色的外边框线"""
    tbl = table._tbl
    ptr = tbl.get_or_add_tblPr()
    borders = OxmlElement('w:tblBorders')
    
    # 蓝色外边框 (1.5pt = sz 12, 颜色 4472C4)
    for border_name in ['top', 'bottom', 'left', 'right']:
        edge = OxmlElement(f'w:{border_name}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), '12')
        edge.set(qn('w:color'), '4472C4')
        borders.append(edge)
    
    # 内部横线 (0.5pt, 蓝色)
    inside_h = OxmlElement('w:insideH')
    inside_h.set(qn('w:val'), 'single')
    inside_h.set(qn('w:sz'), '4')
    inside_h.set(qn('w:color'), '4472C4')
    borders.append(inside_h)
    
    ptr.append(borders)

def draw_custom_pie(ax, values, labels, fig_num=0):
    """饼图：彻底解决数值重叠问题"""
    colors = ['#4472C4', '#ED7D31', '#A5A5A5', '#FFC000', '#5B9BD5', '#70AD47']
    wedges, _ = ax.pie(values, colors=colors[:len(values)], startangle=90, counterclock=False, 
                       wedgeprops={'edgecolor': 'white', 'linewidth': 1})
    
    total = sum(values)
    for i, p in enumerate(wedges):
        ang = (p.theta2 - p.theta1)/2. + p.theta1
        y = np.sin(np.deg2rad(ang))
        x = np.cos(np.deg2rad(ang))
        dist = 1.7
        
        y_text = 1.35 * y
        if abs(y) < 0.3: y_text = 1.5 * y
        x_text = dist * np.sign(x) if x != 0 else dist
        ha = "left" if x_text > 0 else "right"
        
        # 针对图5，强制A向右，B向左延伸
        if fig_num == 5:
            lbl_str = str(labels[i]).strip()
            if lbl_str == 'A':
                x_text = dist
                ha = "left"
                y_text = 1.6
            elif lbl_str == 'B':
                x_text = -dist
                ha = "right"
                y_text = 1.6

        connectionstyle = f"angle,angleA=0,angleB={ang}"
        label_text = f"{labels[i]}\n{int(values[i])} ({(values[i]/total*100):.1f}%)"
        ax.annotate(label_text, xy=(x, y), xytext=(x_text, y_text),
                    horizontalalignment=ha, verticalalignment="center",
                    arrowprops=dict(arrowstyle="-", color="black", connectionstyle=connectionstyle),
                    fontsize=10, fontweight='bold')

def add_page_number(doc):
    for sec in doc.sections:
        footer = sec.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.extend([fldChar1, instrText, fldChar2])

def process_smart(doc, text):
    text = text.replace('\\%', '%').replace('$$', '').replace('\r', '')
    text = re.sub(r'\([\d\.\+\-\*/\s]+\s*[≈=]\s*[\d\.\%]+\)', '', text)
    lines = text.split('\n')
    table_rows, current_title, is_chart_mode = [], "", False

    def flush_table():
        nonlocal table_rows, current_title, is_chart_mode
        if not table_rows: return
        
        # 修改：采用切片方式解析，保留空单元格，确保列数对齐
        raw_data = []
        for r in table_rows:
            if '|' in r and '---' not in r:
                cells = [c.strip() for c in r.split('|')]
                if len(cells) >= 3:
                    raw_data.append(cells[1:-1])
        
        if len(raw_data) >= 2:
            try:
                # 提取表/图编号
                num_match = re.search(r'\d+', current_title)
                current_num = int(num_match.group()) if num_match else 0
                
                if is_chart_mode:
                    df = pd.DataFrame(raw_data[1:], columns=raw_data[0])
                    fig, ax = plt.subplots(figsize=(10, 6))
                    def clean(v):
                        s = re.sub(r'[^\d.]', '', str(v))
                        return float(s) if s else 0.0

                    # 1. 饼图 (4, 5)
                    if current_num in [4, 5]:
                        v_list = [clean(v) for v in df.iloc[:, 1]]
                        draw_custom_pie(ax, v_list, df.iloc[:, 0].tolist(), current_num)
                        ax.set_xlim(-2.5, 2.5); ax.set_ylim(-1.6, 1.6)

                    # 2. 多系列分组柱状图 (2, 12)
                    elif current_num in [2, 12]:
                        custom_bar_colors = ['#F4CCCC', '#F9CB9C', '#FFE599', '#B6D7A8', '#A2C4C9', '#A4C2F4']
                        df_plot = df[~df.iloc[:, 0].str.contains('合计|总计', na=False)]
                        x_labels = [re.sub(r'（.*?）|\(.*?\)|万元|项', '', str(x)) for x in df_plot.iloc[:, 0]]
                        categories = [c for c in df_plot.columns[1:] if '合计' not in c and c.strip()]
                        x = np.arange(len(x_labels))
                        width = 0.8 / (len(categories) + 1)
                        for i, cat in enumerate(categories):
                            vals = [clean(v) for v in df_plot[cat]]
                            ax.bar(x + i*width, vals, width, label=cat, color=custom_bar_colors[i % len(custom_bar_colors)])
                        ax.set_xticks(x + width*(len(categories)-1)/2); ax.set_xticklabels(x_labels)
                        ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.15), ncol=len(categories))
                        plt.subplots_adjust(bottom=0.2)

                    # 3. 双轴图 (9)
                    elif current_num == 9:
                        years = [re.sub(r'（.*?）|\(.*?\)', '', str(x)) for x in df.iloc[:, 0]]
                        counts = [clean(v) for v in df.iloc[:, 1]]; fundings = [clean(v) for v in df.iloc[:, 2]]
                        bars = ax.bar(years, counts, color='#4472C4', label='立项数(项)', width=0.5)
                        for bar in bars:
                            h = bar.get_height()
                            if h > 0: ax.text(bar.get_x() + bar.get_width()/2, h, f'{int(h)}', ha='center', va='bottom', fontsize=10)
                        ax2 = ax.twinx()
                        ax2.plot(years, fundings, color='#ED7D31', marker='o', linewidth=2, label='到账经费(万元)')
                        for i, val in enumerate(fundings):
                            year_str = str(years[i]).strip()
                            va_val = 'top' if '2024' in year_str else 'bottom'
                            ax2.text(years[i], val, f'{val:.2f}', ha='center', va=va_val, fontsize=10, color='#C55A11', 
                                     bbox=dict(facecolor='white', edgecolor='none', alpha=0.6, pad=0.5))
                        fig.legend(loc='lower center', bbox_to_anchor=(0.5, 0.02), ncol=2)
                        plt.subplots_adjust(top=0.88, bottom=0.15)

                    # 4. 横向趋势图 (10, 11)
                    elif current_num in [10, 11]:
                        headers = [re.sub(r'[^\d]', '', str(h)) for h in df.columns]
                        year_headers = [h for h in headers if len(h) == 4]
                        if len(year_headers) >= 3:
                            x_labels = year_headers
                            y_vals = [clean(df.iloc[0, i]) for i in range(1, len(df.columns))]
                        else:
                            x_labels = [re.sub(r'（.*?）|\(.*?\)|万元', '', str(x)) for x in df.iloc[:, 0]]
                            y_vals = [clean(v) for v in df.iloc[:, 1]]
                        bars = ax.bar(x_labels, y_vals, color='#4472C4', width=0.5)
                        for bar in bars:
                            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height(), f'{int(bar.get_height())}', ha='center', va='bottom', fontweight='bold')

                    # 5. 普通柱状图
                    else:
                        x_labels = [str(x).replace('万元', '').strip() for x in df.iloc[:, 0]]
                        vals = [clean(v) for v in df.iloc[:, 1]]
                        bars = ax.bar(x_labels, vals, color='#4472C4', width=0.5)
                        for i, v in enumerate(vals): ax.text(i, v, f'{int(v)}', ha='center', va='bottom')

                    xy_labels = {1:('发表年份','论文数量'), 2:('年份','论文数量'), 3:('年份','立项数量'), 6:('经费区间（万元）','项目数量'), 
                                 7:('立项年份','项目数量'), 8:('经费区间','项目数量'), 9:('年份','项目数量'), 10:('年份','著作出版数量'), 
                                 11:('年份','获奖数量'), 12:('年份','获奖数量')}
                    if current_num in xy_labels:
                        ax.set_xlabel(xy_labels[current_num][0], fontweight='bold')
                        ax.set_ylabel(xy_labels[current_num][1], fontweight='bold')

                    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight'); plt.close(); buf.seek(0)
                    doc.add_picture(buf, width=Inches(5.6))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    # ================= 表格逻辑开始 =================
                    
                    # (A) 表 11, 13, 14 的通用处理：根据数据量动态双栏布局
                    if current_num in [11, 13, 14]:
                        # 设阈值为 11 (1行表头 + 10行数据)，如果总行数超过此值，则平分数据进行双栏显示
                        if len(raw_data) > 11:
                            header = raw_data[0]
                            data_body = raw_data[1:]
                            n = len(data_body)
                            half = (n + 1) // 2
                            
                            new_raw_data = [header + header] # 拼接左右双表头
                            for i in range(half):
                                row_left = data_body[i]
                                # 右侧如果没数据了，补空字符串
                                row_right = data_body[i + half] if (i + half) < n else [""] * len(header)
                                new_raw_data.append(row_left + row_right)
                            raw_data = new_raw_data

                    # (B) 表 12 的特殊处理：按级别（国家级/省部级）分组并左右排版
                    if current_num == 12:
                        df_12 = pd.DataFrame(raw_data[1:], columns=raw_data[0])
                        new_table_data = [["学者", "所属单位", "项目数量", "学者", "所属单位", "项目数量"]]
                        levels = ["国家级", "省部级"]
                        for lvl in levels:
                            sub_df = df_12[df_12['级别'].str.contains(lvl, na=False)]
                            if not sub_df.empty:
                                suffix = "（大于等于2项）" if lvl == "国家级" else "（大于等于3项）"
                                new_table_data.append([lvl + suffix] * 6)
                                sub_rows = sub_df[['负责人', '单位', '立项数']].values.tolist()
                                half = (len(sub_rows) + 1) // 2
                                for i in range(half):
                                    left = sub_rows[i]
                                    right = sub_rows[i + half] if (i + half) < len(sub_rows) else ["", "", ""]
                                    new_table_data.append(left + right)
                        raw_data = new_table_data

                    # 创建表格
                    table = doc.add_table(rows=len(raw_data), cols=len(raw_data[0]))
                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.allow_autofit = False

                    for i, row_values in enumerate(raw_data):
                        table.rows[i].height = Cm(0.71)
                        row_str = "".join([str(x) for x in row_values])
                        
                        # (D) 特殊行合并逻辑：表 10 的 B/C 级标题，或表 12 的级别标题
                        is_special_10 = ("B级" in row_str or "C级" in row_str) and current_num == 10
                        is_special_12 = ("国家级" in row_str or "省部级" in row_str) and current_num == 12 and len(set(row_values)) == 1
                        
                        if is_special_10 or is_special_12:
                            # 提取显示的文字（第一个非空的单元格）
                            display_text = next((x for x in row_values if x.strip()), "").replace('*', '').strip()
                            merged_cell = table.cell(i, 0).merge(table.cell(i, len(row_values)-1))
                            for p in merged_cell.paragraphs: p.clear()
                            p = merged_cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            set_font(p.add_run(display_text), 11, True)
                            if is_special_10: set_cell_shading(merged_cell, "EBF1DE") # 给表10的标题行加淡绿底色以区分
                        else:
                            for j, val in enumerate(row_values):
                                if j >= len(table.columns): break
                                cell = table.cell(i, j)
                                cell.text = str(val).replace('*','')
                                # 动态宽度调整
                                header_text = str(raw_data[0][j])
                                if "序号" in header_text: cell.width = Cm(1.0)
                                elif any(x in header_text for x in ["姓名", "学者", "负责人"]): cell.width = Cm(1.8)
                                elif any(x in header_text for x in ["单位", "学院", "所属单位"]): cell.width = Cm(4.8)
                                elif any(x in header_text for x in ["立项数", "项目数量", "立项数量"]): cell.width = Cm(1.5)
                                elif "级别" in header_text: cell.width = Cm(2.0)
                                else: cell.width = Cm(2.5)

                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                if i == 0:
                                    set_cell_shading(cell, "4472C4")
                                    set_font(p.runs[0] if p.runs else p.add_run(str(val)), 10, True, "FFFFFF")
                                elif i % 2 == 0:
                                    set_cell_shading(cell, "D9E1F2")
                                    set_font(p.runs[0] if p.runs else p.add_run(str(val)), 10, False)
                                else:
                                    set_font(p.runs[0] if p.runs else p.add_run(str(val)), 10, False)
                    set_table_border(table)
            except Exception as e: print(f"Error processing {current_title}: {e}")
        table_rows, is_chart_mode, current_title = [], False, ""

    for line in lines:
        l = line.strip().replace('<center>', '').replace('</center>', '')
        if not l: continue
        if l.startswith('#'):
            flush_table()
            hash_count = l.count('#')
            if hash_count <= 3:
                p = doc.add_heading('', level=hash_count)
                p.paragraph_format.line_spacing = 1.5
                if "附录" in l: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(l.replace('#', '').strip())
                set_font(run, 14, True)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(l.replace('#', '').strip()); set_font(run, 12, True)
        elif re.match(r'(\*\*?)?(附)?[图表]\s?[\d\-\.]+[:：\s]', l):
            flush_table(); current_title = l.replace('*', '').strip(); is_chart_mode = "图" in current_title
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(current_title); set_font(run, 11, True)
        elif l.startswith('|'): table_rows.append(l)
        else:
            if table_rows: flush_table()
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = Pt(24)
            run = p.add_run(l.replace('**', '')); set_font(run, 12)
    flush_table()

def add_toc(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("目  录"); set_font(run, size=16, bold=True)
    p_toc = doc.add_paragraph(); run_toc = p_toc.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    hint_text = OxmlElement('w:t'); hint_text.text = "（请在此处右键单击，选择“更新域”以生成目录内容）"
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run_toc._r.extend([fldChar1, instrText, fldChar2, hint_text, fldChar3])
    doc.add_page_break()

@app.post("/generate_report_word")
async def generate_report_word(input_data: ReportInput, request: Request):
    try:
        doc = Document()
        element = doc.settings.element
        update_fields = OxmlElement('w:updateFields'); update_fields.set(qn('w:val'), 'true'); element.append(update_fields)
        for _ in range(4): doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("山东师范大学人文社会科学科研成果发展态势分析报告\n（2020-2024）")
        set_font(run, size=22, bold=True); doc.add_page_break()
        add_toc(doc); add_page_number(doc)
        for i in range(1, 9): 
            txt = getattr(input_data, f"ch{i}_text", "")
            if txt and txt.strip(): process_smart(doc, txt)
        if input_data.appendix and input_data.appendix.strip():
            doc.add_page_break(); process_smart(doc, input_data.appendix)
        fname = f"report_{uuid.uuid4().hex[:8]}.docx"
        full_path = os.path.join("static", fname)
        doc.save(full_path)
        return {"file": f"{str(request.base_url).rstrip('/')}/static/{fname}", "status": "success"}
    except Exception as e: return {"file": "", "status": "error", "message": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)
